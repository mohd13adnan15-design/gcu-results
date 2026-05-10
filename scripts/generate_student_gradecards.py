#!/usr/bin/env python3
"""
Grade Card Generation Script - Garden City University
Generates personalized grade cards in official GCU format
Matches the official grade card template exactly
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Try to import Supabase, but allow demo mode if not available
try:
    from supabase import create_client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False

# Configuration
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://wmanjhavutkrjbuiwlus.supabase.co")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", "")
OUTPUT_DIR = Path(__file__).parent.parent / "public" / "templates" / "student_gradecards"

# Ensure output directory exists
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_border(cell, **kwargs):
    """Add borders to table cells"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


def add_logo_and_header(doc, student_data):
    """Add GCU logo and university header exactly as in template"""
    # Header table with QR code, logo, student photo placeholders (3 columns)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    # Set table width to full page
    header_table.width = Inches(7.5)
    
    # Left cell - Registration ID and QR code
    left_cell = header_table.rows[0].cells[0]
    left_cell.width = Inches(1.8)
    
    # Add registration at top
    reg_para = left_cell.paragraphs[0]
    reg_para.text = student_data.get("registration_no", "GCUBTRE148")
    reg_run = reg_para.runs[0]
    reg_run.font.size = Pt(10)
    reg_run.font.bold = True
    
    # Add GCUBCA text
    gcubca_para = left_cell.add_paragraph()
    gcubca_para.text = "GCUBCA"
    gcubca_run = gcubca_para.runs[0]
    gcubca_run.font.size = Pt(9)
    gcubca_run.font.bold = True
    
    # Add QR code placeholder
    qr_para = left_cell.add_paragraph()
    qr_para.text = "[QR CODE]"
    qr_run = qr_para.runs[0]
    qr_run.font.size = Pt(8)
    qr_run.font.color.rgb = RGBColor(100, 100, 100)
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Middle cell - Logo and University name
    middle_cell = header_table.rows[0].cells[1]
    middle_cell.width = Inches(3.9)
    middle_cell.vertical_alignment = 1  # Center
    
    # Clear default paragraph
    middle_cell.paragraphs[0].clear()
    
    # Add logo with G in a circle (represented as text)
    logo_para = middle_cell.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run("●")
    logo_run.font.size = Pt(36)
    logo_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red/maroon
    
    # University name
    uni_para = middle_cell.add_paragraph()
    uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uni_run = uni_para.add_run("GARDEN CITY\nUNIVERSITY")
    uni_run.font.size = Pt(16)
    uni_run.font.bold = True
    uni_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red to match screenshot
    
    # Tagline
    tag_para = middle_cell.add_paragraph()
    tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_para.add_run("EMPHASIS ON LIFE")
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Right cell - Student Photo placeholder
    right_cell = header_table.rows[0].cells[2]
    right_cell.width = Inches(1.8)
    right_cell.vertical_alignment = 1
    
    right_para = right_cell.paragraphs[0]
    right_para.text = "Student Photo"
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_run = right_para.runs[0]
    right_run.font.size = Pt(9)
    right_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add borders to header table
    for row in header_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), '12')
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), '000000')
                tcBorders.append(edge_el)
            tcPr.append(tcBorders)
    
    # Add school name in dark red
    school_para = doc.add_paragraph()
    school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_run = school_para.add_run("SCHOOL OF ENGINEERING AND TECHNOLOGY")
    school_run.font.size = Pt(12)
    school_run.font.bold = True
    school_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red


def add_student_info(doc, student_data):
    """Add student information section in table format"""
    # Info table - 4 columns for the layout
    info_table = doc.add_table(rows=3, cols=4)
    info_table.autofit = False
    info_table.width = Inches(7.5)
    
    # Set column widths
    info_table.columns[0].width = Inches(1.8)
    info_table.columns[1].width = Inches(1.9)
    info_table.columns[2].width = Inches(1.8)
    info_table.columns[3].width = Inches(1.8)
    
    # Row 1
    info_table.rows[0].cells[0].text = "PROGRAMME TITLE"
    info_table.rows[0].cells[1].text = student_data.get("programme_title", "Bachelor of Technology in Robotics and Automation")
    info_table.rows[0].cells[2].text = "PROGRAMME CODE:"
    info_table.rows[0].cells[3].text = student_data.get("programme_code", "BTECH")
    
    # Row 2
    info_table.rows[1].cells[0].text = "NAME OF THE STUDENT"
    info_table.rows[1].cells[1].text = student_data.get("name", "Mohammad Anwar Attar")
    info_table.rows[1].cells[2].text = "REGISTRATION NO"
    info_table.rows[1].cells[3].text = student_data.get("registration_no", "24BTRE148")
    
    # Row 3
    info_table.rows[2].cells[0].text = "SEMESTER"
    info_table.rows[2].cells[1].text = student_data.get("semester", "5")
    info_table.rows[2].cells[2].text = "MONTH & YEAR OF THE EXAMINATION"
    info_table.rows[2].cells[3].text = student_data.get("exam_month_year", "November - 2025")
    
    # Add borders to info table
    for row in info_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), '12')
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), '000000')
                tcBorders.append(edge_el)
            tcPr.append(tcBorders)
    
    # Make header cells bold
    for i in range(4):
        for paragraph in info_table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        for paragraph in info_table.rows[1].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        for paragraph in info_table.rows[2].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def add_grades_table(doc, grades_data):
    """Add the main grades table with course categories exactly as in template"""
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("GRADE CARD")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(139, 0, 0)
    
    # Create grades table - 8 columns for all fields
    grades_table = doc.add_table(rows=1, cols=8)
    grades_table.autofit = False
    grades_table.width = Inches(7.5)
    
    # Header row
    header_cells = grades_table.rows[0].cells
    headers = ["SL.\nNo.", "COURSE\nCODE", "COURSE TITLE", "COURSE\nCREDITS", "CREDITS\nEARNED", "GRADE\nOBTAINED", "GRADE\nPOINTS", ""]
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header and center
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    row_counter = 1
    
    # Add CORE COURSE section
    core_courses = grades_data.get("core_courses", [])
    if core_courses:
        # Add category header
        row = grades_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = "CORE COURSE"
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for grade in core_courses:
            row = grades_table.add_row()
            row.cells[0].text = str(row_counter)
            row.cells[1].text = grade.get("course_code", "")
            row.cells[2].text = grade.get("course_title", "")
            row.cells[3].text = str(grade.get("credits", ""))
            row.cells[4].text = str(grade.get("credits_earned", ""))
            row.cells[5].text = grade.get("grade_obtained", "")
            row.cells[6].text = str(grade.get("grade_points", ""))
            row_counter += 1
    
    # Add PRACTICAL section
    practical_courses = grades_data.get("practical_courses", [])
    if practical_courses:
        row = grades_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = "PRACTICAL"
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for grade in practical_courses:
            row = grades_table.add_row()
            row.cells[0].text = str(row_counter)
            row.cells[1].text = grade.get("course_code", "")
            row.cells[2].text = grade.get("course_title", "")
            row.cells[3].text = str(grade.get("credits", ""))
            row.cells[4].text = str(grade.get("credits_earned", ""))
            row.cells[5].text = grade.get("grade_obtained", "")
            row.cells[6].text = str(grade.get("grade_points", ""))
            row_counter += 1
    
    # Add ABILITY ENHANCEMENT COMPULSORY COURSE section
    ability_courses = grades_data.get("ability_courses", [])
    if ability_courses:
        row = grades_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = "ABILITY ENHANCEMENT COMPULSORY COURSE"
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for grade in ability_courses:
            row = grades_table.add_row()
            row.cells[0].text = str(row_counter)
            row.cells[1].text = grade.get("course_code", "")
            row.cells[2].text = grade.get("course_title", "")
            row.cells[3].text = str(grade.get("credits", ""))
            row.cells[4].text = str(grade.get("credits_earned", ""))
            row.cells[5].text = grade.get("grade_obtained", "")
            row.cells[6].text = str(grade.get("grade_points", ""))
            row_counter += 1
    
    # Add SKILL ENHANCEMENT COURSE section
    skill_courses = grades_data.get("skill_courses", [])
    if skill_courses:
        row = grades_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = "SKILL ENHANCEMENT COURSE"
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for grade in skill_courses:
            row = grades_table.add_row()
            row.cells[0].text = str(row_counter)
            row.cells[1].text = grade.get("course_code", "")
            row.cells[2].text = grade.get("course_title", "")
            row.cells[3].text = str(grade.get("credits", ""))
            row.cells[4].text = str(grade.get("credits_earned", ""))
            row.cells[5].text = grade.get("grade_obtained", "")
            row.cells[6].text = str(grade.get("grade_points", ""))
            row_counter += 1
    
    # Add OPEN ELECTIVE COURSE section
    open_elective = grades_data.get("open_elective_courses", [])
    if open_elective:
        row = grades_table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        row.cells[2].text = "OPEN ELECTIVE COURSE"
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for grade in open_elective:
            row = grades_table.add_row()
            row.cells[0].text = str(row_counter)
            row.cells[1].text = grade.get("course_code", "")
            row.cells[2].text = grade.get("course_title", "")
            row.cells[3].text = str(grade.get("credits", ""))
            row.cells[4].text = str(grade.get("credits_earned", ""))
            row.cells[5].text = grade.get("grade_obtained", "")
            row.cells[6].text = str(grade.get("grade_points", ""))
            row_counter += 1
    
    # Add TOTAL row
    row = grades_table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = ""
    row.cells[2].text = "TOTAL"
    row.cells[2].paragraphs[0].runs[0].font.bold = True
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[3].text = str(grades_data.get("total_credits", ""))
    row.cells[4].text = str(grades_data.get("total_credits_earned", ""))
    row.cells[5].text = ""
    row.cells[6].text = ""
    
    # Add borders to grades table
    for row in grades_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), '8')
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), '000000')
                tcBorders.append(edge_el)
            tcPr.append(tcBorders)


def add_footer(doc, student_data):
    """Add footer with date, seal, and signature"""
    # Date
    date_para = doc.add_paragraph()
    date_para.text = f"Date : {datetime.now().strftime('%d-%m-%Y')}"
    
    # Footer table with seal and signature
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.autofit = False
    
    # Left cell - Seal placeholder
    left_cell = footer_table.rows[0].cells[0]
    seal_para = left_cell.paragraphs[0]
    seal_para.text = "GCU SEAL"
    seal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal_run = seal_para.runs[0]
    seal_run.font.size = Pt(10)
    
    # Right cell - Signature line
    right_cell = footer_table.rows[0].cells[1]
    sig_para = right_cell.paragraphs[0]
    sig_para.text = "Controller of Examinations"
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = sig_para.runs[0]
    sig_run.font.size = Pt(10)


def build_grade_card(student_data, grades_data):
    """Build complete grade card document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
    
    # Add components
    add_logo_and_header(doc, student_data)
    add_student_info(doc, student_data)
    add_grades_table(doc, grades_data)
    add_footer(doc, student_data)
    
    return doc


def get_demo_data():
    """Return demo student and grades data matching screenshot format"""
    student_data = {
        "id": "24BTRE148",
        "name": "Mohammad Anwar Attar",
        "programme_title": "Bachelor of Technology in Robotics and Automation",
        "programme_code": "BTECH",
        "registration_no": "24BTRE148",
        "semester": "5",
        "exam_month_year": "November - 2025",
    }
    
    grades_data = {
        "core_courses": [
            {
                "course_code": "24BTSE3512",
                "course_title": "EMBEDDED SYSTEMS",
                "credits": 4.0,
                "credits_earned": 4.0,
                "grade_obtained": "A",
                "grade_points": 8.0
            },
            {
                "course_code": "",
                "course_title": "",
                "credits": 4.0,
                "credits_earned": 4.0,
                "grade_obtained": "A+",
                "grade_points": 9.0
            }
        ],
        "practical_courses": [
            {
                "course_code": "24BTPE3514",
                "course_title": "ROBOT VISION LAB",
                "credits": 2.0,
                "credits_earned": 2.0,
                "grade_obtained": "O",
                "grade_points": 10.0
            },
            {
                "course_code": "24BTPE3515",
                "course_title": "HYDRAULICS AND PNEUMATICS LAB",
                "credits": 2.0,
                "credits_earned": 2.0,
                "grade_obtained": "A",
                "grade_points": 8.0
            }
        ],
        "ability_courses": [
            {
                "course_code": "24AAECC3316",
                "course_title": "BUSINESS COMMUNICATION",
                "credits": 3.0,
                "credits_earned": 3.0,
                "grade_obtained": "A+",
                "grade_points": 9.0
            },
            {
                "course_code": "24AAECC3317",
                "course_title": "INDIAN KNOWLEDGE SYSTEM",
                "credits": 2.0,
                "credits_earned": 2.0,
                "grade_obtained": "B+",
                "grade_points": 7.0
            }
        ],
        "skill_courses": [
            {
                "course_code": "24BTSE3518",
                "course_title": "MACHINE LEARNING ESSENTIALS",
                "credits": 2.0,
                "credits_earned": 2.0,
                "grade_obtained": "A",
                "grade_points": 8.0
            }
        ],
        "open_elective_courses": [
            {
                "course_code": "24AOPEL3521",
                "course_title": "PRINCIPLES OF ECONOMICS",
                "credits": 1.0,
                "credits_earned": 1.0,
                "grade_obtained": "A",
                "grade_points": 8.0
            }
        ],
        "total_credits": 26.0,
        "total_credits_earned": 25.0,
        "final_grade": "A+",
    }
    
    return student_data, grades_data


def fetch_student_data():
    """Fetch student data from Supabase or return demo data"""
    if not SUPABASE_AVAILABLE:
        print("⚠️  Supabase not available. Using demo mode.")
        student_data, grades_data = get_demo_data()
        return [student_data], [grades_data]
    
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
        response = supabase.table("students").select("*").execute()
        return response.data
    except Exception as e:
        print(f"⚠️  Failed to fetch from Supabase: {e}")
        print("Using demo mode instead...")
        student_data, grades_data = get_demo_data()
        return [student_data], [grades_data]


def fetch_grades_for_student(student_id):
    """Fetch grades for a specific student"""
    if not SUPABASE_AVAILABLE:
        return get_demo_data()[1]
    
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
        response = supabase.table("grade_card_details").select("*").eq("student_id", student_id).execute()
        
        # Organize grades by category
        grades_data = {
            "core_courses": [],
            "practical_courses": [],
            "ability_courses": [],
            "skill_courses": [],
            "total_credit_points": 0,
            "final_grade": "A",
        }
        
        for grade in response.data:
            course_data = {
                "course_code": grade.get("course_code", ""),
                "course_title": grade.get("course_title", ""),
                "credits": grade.get("credits", 0),
                "credits_earned": grade.get("credits_earned", 0),
                "grade_obtained": grade.get("grade_obtained", ""),
                "grade_points": grade.get("grade_points", 0),
            }
            
            category = grade.get("course_category", "core_courses")
            if category in grades_data:
                grades_data[category].append(course_data)
        
        return grades_data
    except Exception as e:
        print(f"⚠️  Failed to fetch grades: {e}")
        return get_demo_data()[1]


def generate_gradecards_for_all_students():
    """Generate grade cards for all students"""
    print("🎓 Grade Card Generation System")
    print("=" * 50)
    
    students_data = fetch_student_data()
    
    if isinstance(students_data, tuple):
        students_data = students_data[0]
    
    successful = 0
    failed = 0
    
    for student in students_data:
        try:
            student_id = student.get("id") or student.get("student_id")
            name = student.get("name", "Unknown")
            
            print(f"\n📄 Generating grade card for: {name} ({student_id})")
            
            # Fetch grades for this student
            grades_data = fetch_grades_for_student(student_id)
            
            # Build document
            doc = build_grade_card(student, grades_data)
            
            # Save to file
            output_file = OUTPUT_DIR / f"gradecard_{student_id}.docx"
            doc.save(str(output_file))
            
            print(f"   ✅ Saved to: {output_file}")
            successful += 1
        
        except Exception as e:
            print(f"   ❌ Error: {e}")
            failed += 1
    
    # Summary
    print("\n" + "=" * 50)
    print(f"✅ Successful: {successful}")
    print(f"❌ Failed: {failed}")
    print(f"📁 Output directory: {OUTPUT_DIR}")
    print("=" * 50)


def main():
    """Main entry point"""
    if len(sys.argv) > 1 and sys.argv[1] == "--single":
        # Generate for single demo student
        print("🎓 Grade Card Generation - Demo Mode")
        print("=" * 50)
        
        student_data, grades_data = get_demo_data()
        
        print(f"\n📄 Generating grade card for: {student_data['name']}")
        
        doc = build_grade_card(student_data, grades_data)
        output_file = OUTPUT_DIR / f"gradecard_{student_data['id']}_demo.docx"
        doc.save(str(output_file))
        
        print(f"✅ Grade card saved to: {output_file}")
        print(f"📊 File size: {output_file.stat().st_size / 1024:.1f} KB")
    else:
        # Generate for all students
        generate_gradecards_for_all_students()


if __name__ == "__main__":
    main()
